"""Extract text from PDF, DOCX and TXT attachments. Failures are reported,
never hidden: the caller decides whether the submission has usable text."""
import logging
import re
from pathlib import Path

log = logging.getLogger("extract")
SUPPORTED = {".pdf", ".docx", ".txt"}


def normalize(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+\n", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def extract_pdf(path: Path) -> str:
    from pypdf import PdfReader
    reader = PdfReader(str(path))
    parts = []
    for page in reader.pages:
        parts.append(page.extract_text() or "")
    return normalize("\n\n".join(parts))


def extract_docx(path: Path) -> str:
    import docx
    document = docx.Document(str(path))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return normalize("\n".join(parts))


def extract_txt(path: Path) -> str:
    raw = path.read_bytes()
    for enc in ("utf-8", "utf-16", "latin-1"):
        try:
            return normalize(raw.decode(enc))
        except UnicodeDecodeError:
            continue
    return ""


def extract(path) -> dict:
    """Return {"status": "ok"|"unsupported"|"error", "text": str, "error": str|None}."""
    path = Path(path)
    ext = path.suffix.lower()
    if ext not in SUPPORTED:
        return {"status": "unsupported", "text": "", "error": f"unsupported format {ext}"}
    try:
        if ext == ".pdf":
            text = extract_pdf(path)
        elif ext == ".docx":
            text = extract_docx(path)
        else:
            text = extract_txt(path)
    except Exception as exc:  # noqa: BLE001
        log.warning("extraction failed for %s: %s", path.name, exc)
        return {"status": "error", "text": "", "error": str(exc)}
    if not text.strip():
        return {"status": "error", "text": "", "error": "no extractable text (scanned image?)"}
    return {"status": "ok", "text": text, "error": None}
